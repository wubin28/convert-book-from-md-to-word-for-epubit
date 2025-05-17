import sys
import os
import re
import shutil
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml import OxmlElement

def read_markdown_file(file_path):
    """Read content from a markdown file."""
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def copy_template(template_path, output_path):
    """Copy the template DOCX file to a new file."""
    shutil.copy(template_path, output_path)

def get_style_name(doc, default_name, style_type=WD_STYLE_TYPE.PARAGRAPH):
    """Try to find an appropriate style name in the document."""
    # Special handling for list styles which are often stored as paragraph styles in templates
    if default_name in ['List Bullet', 'List Number'] and style_type == WD_STYLE_TYPE.PARAGRAPH:
        # Just return the name, we'll handle it specially when applying
        return default_name
    
    # Check if the exact style exists
    try:
        doc.styles.get_style_id(default_name, style_type)
        return default_name
    except KeyError:
        pass
    except ValueError:
        # Style exists but wrong type, just return None
        return None
    
    # Common variations for English and other languages
    variations = {
        'Heading1': ['Heading 1', '标题 1', '标题1', 'Heading1', '1', 'Title'],
        'Heading2': ['Heading 2', '标题 2', '标题2', 'Heading2', '2'],
        'Heading3': ['Heading 3', '标题 3', '标题3', 'Heading3', '3'],
        'Heading4': ['Heading 4', '标题 4', '标题4', 'Heading4', '4'],
        'Caption': ['Caption', '标题', '图表标题', 'Table Caption', 'Figure Caption'],
        'List Bullet': ['List Bullet', '项目符号', 'Bulleted List', '无序列表'],
        'List Number': ['List Number', '编号', 'Numbered List', '有序列表'],
        'Normal': ['Normal', '正文', 'Body Text', 'Regular'],
        'No Spacing': ['No Spacing', '无间距', 'Compact'],
        'Intense Quote': ['Intense Quote', '强调引用', 'Strong Quote'],
        'Quote': ['Quote', '引用', 'Block Quote']
    }
    
    # Try to find a matching style
    if default_name in variations:
        for style_name in variations[default_name]:
            try:
                doc.styles.get_style_id(style_name, style_type)
                return style_name
            except (KeyError, ValueError):
                continue
    
    # Fallback to Normal style for paragraphs
    if style_type == WD_STYLE_TYPE.PARAGRAPH:
        try:
            doc.styles.get_style_id('Normal', style_type)
            return 'Normal'
        except (KeyError, ValueError):
            pass
    
    # If nothing works, return None and handle it later
    return None

def find_table_style(doc):
    """Find a suitable table style."""
    table_styles = ['Table Grid', 'Grid Table 1 Light', 'Plain Table 1', 'Table Normal']
    for style_name in table_styles:
        try:
            # Just check if the style exists
            if style_name in doc.styles:
                return style_name
        except:
            continue
    return None

def add_gray_background(paragraph):
    """Add light gray background to a paragraph."""
    try:
        # Use light gray color (F5F5F5 is a very light gray)
        shading_xml = f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>'
        shading_element = OxmlElement(shading_xml)
        paragraph._element.get_or_add_pPr().append(shading_element)
    except:
        # If there's any error, just continue without adding background
        pass

def set_cell_background(cell, color):
    """Set cell background color using direct XML manipulation."""
    # Create the shading element with the specified fill color
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)  # Set fill color
    shading.set(qn('w:val'), 'clear')  # Clear any previous shading
    
    # Get or create the cell properties element
    tcPr = cell._tc.get_or_add_tcPr()
    
    # Remove any existing shading
    for s in tcPr.findall(qn('w:shd')):
        tcPr.remove(s)
    
    # Add the new shading
    tcPr.append(shading)

def add_bullet_list_item(doc, text, list_style=None):
    """
    Add a bullet list item with proper formatting using Word's native bullet points.
    Properly handles bold text marked with **.
    """
    p = doc.add_paragraph()
    
    # Try to apply built-in bullet list style
    try:
        if list_style:
            p.style = list_style
        else:
            # Use a style that's likely to be in most Word documents
            p.style = 'List Bullet'
    except:
        # If built-in style fails, fall back to manual formatting
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
    
    # Process bold markers in the text
    segments = process_bold_text(text)
    
    # Add text segments with appropriate formatting
    for segment_text, is_bold in segments:
        run = p.add_run(segment_text)
        run.bold = is_bold
    
    return p

def add_manual_numbered_item(doc, text, number=None):
    """
    Add a numbered list item with manual formatting, preserving the original (X) format.
    Properly handles bold text marked with **.
    """
    p = doc.add_paragraph()
    
    # Apply manual formatting with left indentation
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    
    # Process bold markers in the text
    segments = process_bold_text(text)
    
    # Create the prefix with the original number format
    prefix = f"({number}) " if number else ""
    
    # Add the first segment with the prefix
    if segments:
        first_segment = segments[0]
        if first_segment[1]:  # If the first segment should be bold
            run = p.add_run(prefix + first_segment[0])
            run.bold = True
        else:
            run = p.add_run(prefix + first_segment[0])
        
        # Add remaining segments
        for segment_text, is_bold in segments[1:]:
            run = p.add_run(segment_text)
            run.bold = is_bold
    
    return p

def process_bold_text(text):
    """
    Process text to identify bold sections marked with ** in Markdown.
    Returns a list of tuples (text_segment, is_bold).
    """
    segments = []
    
    # Pattern to match **bold text**
    pattern = r'\*\*(.*?)\*\*'
    
    # Find all occurrences of bold text
    bold_matches = list(re.finditer(pattern, text))
    
    if not bold_matches:
        # No bold text found, return the original text as not bold
        return [(text, False)]
    
    # Process text with bold sections
    last_end = 0
    for match in bold_matches:
        # Add the text before this bold section (if any)
        if match.start() > last_end:
            segments.append((text[last_end:match.start()], False))
        
        # Add the bold text without ** markers
        bold_text = match.group(1)
        segments.append((bold_text, True))
        
        last_end = match.end()
    
    # Add any remaining text after the last bold section
    if last_end < len(text):
        segments.append((text[last_end:], False))
    
    return segments

def process_code_block(doc, lines, i, styles):
    """
    Process a markdown code block and add it to the document with proper formatting.
    
    Args:
        doc: The Word document object
        lines: All lines of the markdown content
        i: Current line index
        styles: Document styles dictionary
        
    Returns:
        Updated line index after processing the code block
    """
    # Get language if specified (e.g., ```markdown, ```python)
    code_lang = lines[i][3:].strip()
    
    # Move to the first line after the opening ```
    i += 1
    code_lines = []
    
    # Collect all lines until the closing ```
    while i < len(lines) and not lines[i].startswith('```'):
        code_lines.append(lines[i])
        i += 1
    
    # Skip the closing ``` marker
    i += 1
    
    # Check for code listing title (特定的中文标题格式)
    if i < len(lines) and re.match(r'^代码清单\d+-\d+', lines[i]):
        # Use add_paragraph_with_formatting to handle bold text in code listing title
        p = add_paragraph_with_formatting(doc, lines[i], styles['Caption'])
        i += 1
    
    # Add each line of the code block
    for code_line in code_lines:
        p = doc.add_paragraph(code_line)
        # Apply "No Spacing" style if available
        if styles['No Spacing']:
            try:
                p.style = styles['No Spacing']
            except:
                pass
        
        # Add light gray background
        add_gray_background(p)
        
        # Format with monospace font
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
    
    return i

def add_paragraph_with_formatting(doc, text, style=None):
    """Add a paragraph with proper formatting for bold text marked with **."""
    p = doc.add_paragraph()
    
    # Apply style if provided
    if style:
        try:
            p.style = style
        except:
            pass
    
    # Process bold markers
    segments = process_bold_text(text)
    
    # Add segments with appropriate formatting
    for segment_text, is_bold in segments:
        run = p.add_run(segment_text)
        run.bold = is_bold
    
    return p

def create_bidi_box(doc, title_text, content_lines, title_color='E36C09', content_color='FDE9D9'):
    """Create a styled box for the 避坑指南 section with orange background."""
    # Create the title paragraph with dark orange background
    p_title = doc.add_paragraph()
    p_title._element.get_or_add_pPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{title_color}" w:val="clear"/>')
    )
    
    # Add the title text (make it bold and WHITE)
    run = p_title.add_run(title_text)
    run.bold = True
    # Set the font color to white
    run.font.color.rgb = RGBColor(255, 255, 255)  # RGB value for white
    
    # Add the content with proper formatting for bold text
    for line in content_lines:
        if line.strip():
            # Create paragraph with light orange background
            p_content = doc.add_paragraph()
            p_content._element.get_or_add_pPr().append(
                parse_xml(f'<w:shd {nsdecls("w")} w:fill="{content_color}" w:val="clear"/>')
            )
            
            # Process and add text with bold formatting
            segments = process_bold_text(line)
            for segment_text, is_bold in segments:
                run = p_content.add_run(segment_text)
                run.bold = is_bold

def process_aside_block(doc, lines, i):
    """
    Process aside sections (【避坑指南】) in markdown and format them with orange backgrounds.
    
    Args:
        doc: The Word document object
        lines: All lines of the markdown content
        i: Current line index
        
    Returns:
        Updated line index after processing the aside block
    """
    line = lines[i]
    
    # Handle <aside> tag followed by 【避坑指南】
    if line.strip() == "<aside>":
        i += 1  # Move to next line that should contain 【避坑指南】
        if i < len(lines) and "【避坑指南】" in lines[i]:
            line = lines[i]
        else:
            return i  # Skip malformed aside
    
    # Go through the content to collect all lines
    i += 1  # Move past the header
    content_lines = []
    
    while i < len(lines):
        if lines[i].strip() == "</aside>" or lines[i].strip().startswith("#") or "【避坑指南】" in lines[i]:
            break
        
        if lines[i].strip():
            content_lines.append(lines[i])
        i += 1
    
    # Create the styled box with orange backgrounds
    create_bidi_box(doc, line, content_lines)
    
    # Skip </aside> tag if present
    if i < len(lines) and lines[i].strip() == "</aside>":
        i += 1
    
    return i

def convert_markdown_to_docx(markdown_content, template_path, output_path, markdown_file=None):
    """Convert markdown content to DOCX using the template as a base."""
    # Copy template to output file
    copy_template(template_path, output_path)
    
    # Create a new document from the template
    doc = Document(output_path)
    
    # Save original style definitions
    styles = {}
    for style_type in ['Heading1', 'Heading2', 'Heading3', 'Heading4', 'Caption', 
                      'Normal', 'No Spacing', 'Intense Quote', 'Quote']:
        styles[style_type] = get_style_name(doc, style_type)
    
    # Get list styles
    bullet_list_style = get_style_name(doc, 'List Bullet')
    
    # Find a suitable table style
    table_style = find_table_style(doc)
    
    # Clear existing content but preserve styles
    paragraphs_to_remove = []
    for i, paragraph in enumerate(doc.paragraphs):
        paragraphs_to_remove.append(i)
    
    # Remove paragraphs in reverse order to avoid index issues
    for i in reversed(paragraphs_to_remove):
        if i < len(doc.paragraphs):
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)
    
    # Remove all tables
    for table in doc.tables:
        element = table._element
        element.getparent().remove(element)
    
    # Split content into lines
    lines = markdown_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Process headings
        if line.startswith('# '):  # Chapter title
            p = doc.add_paragraph()
            if styles['Heading1']:
                try:
                    p.style = styles['Heading1']
                except:
                    pass
            
            # Process bold text in headings
            segments = process_bold_text(line[2:])
            for segment_text, is_bold in segments:
                run = p.add_run(segment_text)
                # Only apply bold to segments marked with **
                if is_bold:
                    run.bold = True
            
            i += 1
            continue
        
        if line.startswith('## '):  # Section title
            p = doc.add_paragraph()
            if styles['Heading2']:
                try:
                    p.style = styles['Heading2']
                except:
                    pass
            
            # Process bold text in headings
            segments = process_bold_text(line[3:])
            for segment_text, is_bold in segments:
                run = p.add_run(segment_text)
                # Only apply bold to segments marked with **
                if is_bold:
                    run.bold = True
            
            i += 1
            continue
        
        if line.startswith('### '):  # Subsection title
            p = doc.add_paragraph()
            if styles['Heading3']:
                try:
                    p.style = styles['Heading3']
                except:
                    pass
            
            # Process bold text in headings
            segments = process_bold_text(line[4:])
            for segment_text, is_bold in segments:
                run = p.add_run(segment_text)
                # Only apply bold to segments marked with **
                if is_bold:
                    run.bold = True
            
            i += 1
            continue
        
        if line.startswith('#### '):  # Fourth level title
            p = doc.add_paragraph()
            if styles['Heading4']:
                try:
                    p.style = styles['Heading4']
                except:
                    pass
            
            # Process bold text in headings
            segments = process_bold_text(line[5:])
            for segment_text, is_bold in segments:
                run = p.add_run(segment_text)
                # Only apply bold to segments marked with **
                if is_bold:
                    run.bold = True
            
            i += 1
            continue
        
        # Process code blocks
        if line.startswith('```'):
            i = process_code_block(doc, lines, i, styles)
            continue
        
        # Process aside sections (【避坑指南】)
        if line.strip() == "<aside>" or "【避坑指南】" in line:
            i = process_aside_block(doc, lines, i)
            continue
        
        # Process bullet lists
        if line.strip().startswith('- '):
            items = []
            while i < len(lines) and lines[i].strip().startswith('- '):
                items.append(lines[i].strip()[2:])
                i += 1
            
            for item in items:
                # Use new function to create native Word bullet points
                add_bullet_list_item(doc, item, bullet_list_style)
            
            continue
        
        # Process numbered lists with parentheses (X)
        if re.match(r'^\(\d+\)', line.strip()):
            items = []
            numbers = []
            while i < len(lines) and re.match(r'^\(\d+\)', lines[i].strip()):
                item_match = re.match(r'^\((\d+)\)\s*(.+)$', lines[i].strip())
                if item_match:
                    numbers.append(item_match.group(1))
                    items.append(item_match.group(2))
                i += 1
            
            for idx, item in enumerate(items):
                # Use the manual numbered list function to preserve original numbering
                add_manual_numbered_item(doc, item, numbers[idx])
            
            continue
        
        # Process tables
        if line.strip().startswith('|') and '|' in line[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                if not lines[i].strip().startswith('|-'):  # Skip separator lines
                    table_lines.append(lines[i])
                i += 1
                # Skip separator line
                if i < len(lines) and lines[i].strip().startswith('|-'):
                    i += 1
            
            # Skip empty lines after table
            while i < len(lines) and not lines[i].strip():
                i += 1
            
            # Check for table caption
            if i < len(lines) and lines[i].startswith('>') and '表' in lines[i]:
                # Use add_paragraph_with_formatting to handle bold text in table caption
                p = add_paragraph_with_formatting(doc, lines[i].strip('> '), styles['Caption'])
                i += 1
            
            # Process table content
            if len(table_lines) >= 1:  # Allow tables with just a header
                header_row = table_lines[0]
                data_rows = table_lines[1:] if len(table_lines) > 1 else []
                
                # Parse header row
                header_cells = [cell.strip() for cell in header_row.split('|')[1:-1] if cell.strip()]
                
                if not header_cells:  # Handle malformed table
                    continue
                
                # Create table
                table = doc.add_table(rows=1, cols=len(header_cells))
                if table_style:
                    try:
                        table.style = table_style
                    except:
                        pass
                
                # Add header
                for j, cell_text in enumerate(header_cells):
                    # Process bold text in table header cells
                    cell = table.cell(0, j)
                    p = cell.paragraphs[0]
                    segments = process_bold_text(cell_text)
                    
                    # Clear existing text in the paragraph
                    p.clear()
                    
                    # Add segments with appropriate formatting
                    for segment_text, is_bold in segments:
                        run = p.add_run(segment_text)
                        run.bold = is_bold
                
                # Add data rows
                for row_text in data_rows:
                    cells = [cell.strip() for cell in row_text.split('|')[1:-1]]
                    if not cells:
                        continue
                    
                    row_cells = table.add_row().cells
                    for j, cell_text in enumerate(cells[:len(header_cells)]):
                        # Process bold text in table data cells
                        cell = row_cells[j]
                        p = cell.paragraphs[0]
                        segments = process_bold_text(cell_text)
                        
                        # Clear existing text in the paragraph
                        p.clear()
                        
                        # Add segments with appropriate formatting
                        for segment_text, is_bold in segments:
                            run = p.add_run(segment_text)
                            run.bold = is_bold
            
            continue
        
        # Process image references
        if line.startswith('!['):
            # Extract image information using regex
            match = re.search(r'!\[(.*?)\]\((.*?)\)', line)
            if match:
                alt_text = match.group(1)
                img_path = match.group(2)
                
                # Resolve the image path relative to the current script directory
                # If the image path is relative, assume it's in the same directory
                if not os.path.isabs(img_path):
                    img_path = os.path.join(os.path.dirname(os.path.abspath(markdown_file)), img_path)
                
                try:
                    # Check if image exists
                    if os.path.exists(img_path):
                        # Add image to document with reasonable width
                        doc.add_picture(img_path, width=Inches(6))
                        
                        # Add caption with alt text if provided
                        if alt_text:
                            caption = doc.add_paragraph(alt_text)
                            if styles['Caption']:
                                try:
                                    caption.style = styles['Caption']
                                except:
                                    caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    for run in caption.runs:
                                        run.italic = True
                                        run.font.size = Pt(10)
                    else:
                        # Add placeholder for missing image
                        doc.add_paragraph(f"[Image not found: {img_path}]")
                except Exception as e:
                    # Handle any errors during image processing
                    doc.add_paragraph(f"[Error adding image: {str(e)}]")
            
            i += 1
            continue

        # Process notes (注意, 提示)
        if line.strip() in ["注意", "提示"]:
            p = doc.add_paragraph(line.strip())
            if styles['Intense Quote']:
                try:
                    p.style = styles['Intense Quote']
                except:
                    run = p.runs[0] if p.runs else p.add_run()
                    run.bold = True
            else:
                run = p.runs[0] if p.runs else p.add_run()
                run.bold = True
            i += 1
            
            note_content = []
            while i < len(lines) and lines[i].strip() and not lines[i].startswith('#') and not lines[i].startswith('<'):
                note_content.append(lines[i])
                i += 1
            
            # Create a box for note content
            if note_content:
                table = doc.add_table(rows=1, cols=1)
                if table_style:
                    try:
                        table.style = table_style
                    except:
                        pass
                
                cell = table.cell(0, 0)
                
                for note_line in note_content:
                    # Process bold text in note content
                    p = cell.add_paragraph()
                    segments = process_bold_text(note_line)
                    
                    for segment_text, is_bold in segments:
                        run = p.add_run(segment_text)
                        run.bold = is_bold
                    
                    if styles['Quote']:
                        try:
                            p.style = styles['Quote']
                        except:
                            pass
            
            continue
        
        # Process regular paragraphs
        if line.strip():
            # Use the new function to handle bold text in regular paragraphs
            p = add_paragraph_with_formatting(doc, line, styles['Normal'])
            i += 1
            continue
        
        # Skip empty lines
        i += 1
    
    # Save the document
    doc.save(output_path)
    print(f"Conversion complete! Document saved to {output_path}")

def main():
    if len(sys.argv) != 2:
        print("Usage: python converter.py <markdown_file>")
        sys.exit(1)
    
    markdown_file = sys.argv[1]
    template_path = "to-word-template.docx"
    
    # Generate output path based on the markdown file
    markdown_dir = os.path.dirname(markdown_file)
    markdown_basename = os.path.basename(markdown_file)
    markdown_name = os.path.splitext(markdown_basename)[0]
    output_path = os.path.join(markdown_dir, f"{markdown_name}.docx")
    
    # Check if files exist
    if not os.path.exists(markdown_file):
        print(f"Error: {markdown_file} does not exist")
        sys.exit(1)
    
    if not os.path.exists(template_path):
        print(f"Error: {template_path} does not exist")
        sys.exit(1)
    
    # Read markdown content
    content = read_markdown_file(markdown_file)
    
    # Convert markdown to DOCX
    convert_markdown_to_docx(content, template_path, output_path, markdown_file)    
    print(f"\nNote: The converter has made its best attempt to match the formatting.")
    print(f"      Please open '{output_path}' to verify the result.")
    print(f"      Text wrapped in ** markdown has been converted to bold formatting.")
    print(f"      Code blocks now have a light gray background for better readability.")
    print(f"      Special sections like '【避坑指南】' now have orange backgrounds with white title text as requested.")
    print(f"      Numbered lists keep their original '(1)' formatting.")

if __name__ == "__main__":
    main()